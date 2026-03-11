"""
genera_contratto.py
Genera contratto DOCX per i partner Evolution PRO usando python-docx.
"""
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def genera_contratto(dati: dict, output_path: str) -> str:
    """
    Genera il contratto DOCX con i dati del partner.
    
    Args:
        dati: Dizionario con i dati del partner
        output_path: Path del file output
    
    Returns:
        Path del file generato
    """
    # Assicurati che l'estensione sia .docx
    if not output_path.endswith('.docx'):
        output_path = output_path.replace('.pdf', '.docx').replace('.PDF', '.docx')
        if not output_path.endswith('.docx'):
            output_path += '.docx'
    
    # Crea documento
    doc = Document()
    
    # Stili
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # === INTESTAZIONE ===
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("CONTRATTO DI PARTNERSHIP")
    run.bold = True
    run.font.size = Pt(16)
    
    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = header2.add_run("EVOLUTION PRO")
    run2.bold = True
    run2.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # === PARTI CONTRAENTI ===
    doc.add_paragraph().add_run("TRA").bold = True
    
    doc.add_paragraph(
        "BERTOGLIATTI CLAUDIO, nato a Torino (TO) il 04/11/1974, C.F. BRTCLD74S04L219V, "
        "P.IVA 11788830011, residente in Via Crevacuore 7, 10146 Torino (TO), "
        "di seguito denominato 'EVOLUTION PRO'"
    )
    
    doc.add_paragraph().add_run("E").bold = True
    
    nome = dati.get("nome", "")
    cognome = dati.get("cognome", "")
    azienda = dati.get("azienda", "")
    indirizzo = dati.get("indirizzo", "")
    citta = dati.get("citta", "")
    cap = dati.get("cap", "")
    prov = dati.get("prov", "")
    cf = dati.get("codice_fiscale", "")
    piva = dati.get("partita_iva", "")
    email = dati.get("email", "")
    pec = dati.get("pec", "")
    iban = dati.get("iban", "")
    data_firma = dati.get("data_firma", datetime.now().strftime("%d/%m/%Y"))
    
    partner_text = f"{nome} {cognome}"
    if azienda:
        partner_text += f", titolare di {azienda}"
    partner_text += f", C.F. {cf}"
    if piva:
        partner_text += f", P.IVA {piva}"
    partner_text += f", residente in {indirizzo}, {cap} {citta} ({prov})"
    partner_text += f", email: {email}"
    if pec:
        partner_text += f", PEC: {pec}"
    partner_text += ", di seguito denominato 'PARTNER'"
    
    doc.add_paragraph(partner_text)
    
    doc.add_paragraph()
    
    # === PREMESSE ===
    doc.add_paragraph().add_run("PREMESSO CHE").bold = True
    
    premesse = [
        "Evolution PRO è un'agenzia specializzata nella creazione e lancio di videocorsi online per professionisti;",
        "Il Partner possiede competenze specialistiche nel proprio settore di riferimento;",
        "Le parti intendono collaborare per la creazione, il lancio e la vendita di un videocorso online basato sulle competenze del Partner;",
        "Evolution PRO fornirà supporto strategico, tecnico e operativo per la realizzazione del progetto."
    ]
    
    for p in premesse:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(p)
    
    doc.add_paragraph()
    
    # === ARTICOLI ===
    doc.add_paragraph().add_run("SI CONVIENE E SI STIPULA QUANTO SEGUE").bold = True
    doc.add_paragraph()
    
    # Articolo 1
    art1 = doc.add_paragraph()
    art1.add_run("Art. 1 - OGGETTO DEL CONTRATTO").bold = True
    doc.add_paragraph(
        "Evolution PRO si impegna a fornire al Partner i seguenti servizi:\n"
        "a) Analisi strategica del posizionamento e della nicchia di mercato;\n"
        "b) Supporto nella strutturazione del videocorso;\n"
        "c) Assistenza tecnica per la registrazione e il montaggio dei video;\n"
        "d) Creazione della piattaforma di vendita su Systeme.io;\n"
        "e) Strategia di lancio e marketing;\n"
        "f) Supporto continuativo post-lancio."
    )
    
    # Articolo 2
    doc.add_paragraph()
    art2 = doc.add_paragraph()
    art2.add_run("Art. 2 - INVESTIMENTO INIZIALE").bold = True
    doc.add_paragraph(
        "Il Partner verserà a Evolution PRO un investimento iniziale di € 2.100,00 (duemilacento/00) "
        "IVA inclusa, a copertura del 30% dei costi di produzione.\n\n"
        "Evolution PRO coprirà il restante 70% dei costi di produzione."
    )
    
    # Articolo 3
    doc.add_paragraph()
    art3 = doc.add_paragraph()
    art3.add_run("Art. 3 - ROYALTIES").bold = True
    doc.add_paragraph(
        "Il Partner riconoscerà a Evolution PRO una royalty pari al 10% (dieci per cento) "
        "del fatturato netto derivante dalla vendita del videocorso.\n\n"
        "Le royalties saranno calcolate mensilmente e versate entro il 15 del mese successivo."
    )
    
    # Articolo 4
    doc.add_paragraph()
    art4 = doc.add_paragraph()
    art4.add_run("Art. 4 - PIANO CONTINUITÀ").bold = True
    doc.add_paragraph(
        "Al termine dei primi 12 mesi, il Partner potrà sottoscrivere un Piano Continuità "
        "per mantenere attivo il supporto di Evolution PRO.\n\n"
        "Le condizioni del Piano Continuità saranno comunicate entro il 10° mese di partnership."
    )
    
    # Articolo 5
    doc.add_paragraph()
    art5 = doc.add_paragraph()
    art5.add_run("Art. 5 - PROPRIETÀ INTELLETTUALE").bold = True
    doc.add_paragraph(
        "I contenuti del videocorso restano di proprietà esclusiva del Partner.\n"
        "Evolution PRO mantiene la proprietà dei template, delle strategie e degli strumenti sviluppati."
    )
    
    # Articolo 6
    doc.add_paragraph()
    art6 = doc.add_paragraph()
    art6.add_run("Art. 6 - DURATA").bold = True
    doc.add_paragraph(
        "Il presente contratto ha durata di 12 (dodici) mesi dalla data di sottoscrizione.\n"
        "Alla scadenza, potrà essere rinnovato previo accordo tra le parti."
    )
    
    # Articolo 7
    doc.add_paragraph()
    art7 = doc.add_paragraph()
    art7.add_run("Art. 7 - RECESSO").bold = True
    doc.add_paragraph(
        "Ciascuna parte potrà recedere dal contratto con preavviso scritto di 30 giorni.\n"
        "In caso di recesso anticipato, non è previsto alcun rimborso degli importi già versati."
    )
    
    # Articolo 8
    doc.add_paragraph()
    art8 = doc.add_paragraph()
    art8.add_run("Art. 8 - FORO COMPETENTE").bold = True
    doc.add_paragraph(
        "Per qualsiasi controversia derivante dal presente contratto sarà competente "
        "in via esclusiva il Foro di Torino."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # === DATI BANCARI ===
    doc.add_paragraph().add_run("DATI BANCARI PER IL VERSAMENTO").bold = True
    doc.add_paragraph(
        "Intestatario: BERTOGLIATTI CLAUDIO\n"
        "IBAN: LT94 3250 0974 4929 5781\n"
        "Causale: Partnership Evolution PRO - " + f"{nome} {cognome}"
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # === FIRME ===
    doc.add_paragraph(f"Torino, {data_firma}")
    doc.add_paragraph()
    
    # Tabella firme
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "EVOLUTION PRO"
    table.cell(0, 1).text = "IL PARTNER"
    table.cell(1, 0).text = "\n\n_________________________\nClaudio Bertogliatti"
    table.cell(1, 1).text = f"\n\n_________________________\n{nome} {cognome}"
    
    # Assicura directory output
    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Salva
    doc.save(output_path)
    
    return output_path


if __name__ == "__main__":
    # Leggi input da stdin o da argv
    if len(sys.argv) > 1:
        # Vecchio formato: python genera_contratto.py '{"nome":...}' '/path/output.docx'
        try:
            dati = json.loads(sys.argv[1])
            output = sys.argv[2] if len(sys.argv) > 2 else "/app/backend/static/contratti/output.docx"
        except:
            input_data = json.loads(sys.stdin.read())
            dati = input_data.get("dati", input_data)
            output = input_data.get("output_path", "/app/backend/static/contratti/output.docx")
    else:
        # Nuovo formato: stdin JSON
        input_data = json.loads(sys.stdin.read())
        dati = input_data.get("dati", input_data)
        output = input_data.get("output_path", "/app/backend/static/contratti/output.docx")
    
    # Forza estensione .docx
    output = output.replace('.pdf', '.docx')
    if not output.endswith('.docx'):
        output += '.docx'
    
    result = genera_contratto(dati, output)
    print(f"OK:{result}")
