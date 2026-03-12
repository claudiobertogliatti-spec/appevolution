"""
genera_analisi_docx.py
Genera l'Analisi Strategica Evolution PRO in formato DOCX professionale.

Template ufficiale con:
- Copertina professionale
- Sezioni su pagine separate
- Blocchi risposta cliente con barra gialla
- Insight strategici evidenziati
- Footer su ogni pagina
"""

import os
import sys
import json
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════════════════════════════════════
# COSTANTI - PALETTE EVOLUTION PRO
# ═══════════════════════════════════════════════════════════════════════════════
BLU_SCURO = RGBColor(0x15, 0x1F, 0x28)      # #151F28
GIALLO = RGBColor(0xFD, 0xD3, 0x2A)         # #FDD32A
GRIGIO = RGBColor(0x6B, 0x72, 0x80)         # #6B7280
GRIGIO_CHIARO = RGBColor(0xF9, 0xFA, 0xFB)  # #F9FAFB
NERO = RGBColor(0x00, 0x00, 0x00)


def set_cell_shading(cell, color_hex):
    """Imposta il colore di sfondo di una cella."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_break(doc):
    """Aggiunge un'interruzione di pagina."""
    doc.add_page_break()


def create_yellow_line(doc, width_cm=6):
    """Crea una linea gialla orizzontale usando una tabella."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    cell.width = Cm(width_cm)
    
    # Imposta altezza e colore
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Sfondo giallo
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'FDD32A')
    tcPr.append(shd)
    
    # Imposta altezza della riga
    tr = tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '60')  # circa 2pt
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)
    
    # Rimuovi bordi
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)
    
    cell.paragraphs[0].clear()
    doc.add_paragraph()


def add_section_title(doc, title):
    """Aggiunge un titolo di sezione con linea gialla."""
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BLU_SCURO
    p.space_after = Pt(8)
    
    create_yellow_line(doc)


def add_subtitle(doc, text):
    """Aggiunge un sottotitolo."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLU_SCURO
    p.space_before = Pt(16)
    p.space_after = Pt(8)


def add_body_text(doc, text):
    """Aggiunge testo normale."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = NERO
    p.space_after = Pt(8)


def add_client_response_box(doc, label, response):
    """Aggiunge il blocco risposta cliente con barra gialla laterale."""
    # Label
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GRIGIO
    p.space_after = Pt(4)
    
    # Box con barra gialla (usando tabella)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    
    # Colonna 1: barra gialla
    cell1 = table.cell(0, 0)
    cell1.width = Cm(0.3)
    set_cell_shading(cell1, 'FDD32A')
    cell1.paragraphs[0].clear()
    
    # Colonna 2: contenuto
    cell2 = table.cell(0, 1)
    cell2.width = Cm(15)
    set_cell_shading(cell2, 'F9FAFB')
    
    # Aggiungi padding
    p = cell2.paragraphs[0]
    p.clear()
    run = p.add_run(response)
    run.font.size = Pt(11)
    run.font.color.rgb = NERO
    
    # Spazio dopo la tabella
    doc.add_paragraph()


def add_insight_box(doc, text):
    """Aggiunge il blocco INSIGHT STRATEGICO."""
    # Tabella per simulare il box
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    
    # Colonna 1: barra gialla
    cell1 = table.cell(0, 0)
    cell1.width = Cm(0.4)
    set_cell_shading(cell1, 'FDD32A')
    cell1.paragraphs[0].clear()
    
    # Colonna 2: contenuto
    cell2 = table.cell(0, 1)
    cell2.width = Cm(14)
    
    # Titolo INSIGHT
    p1 = cell2.paragraphs[0]
    p1.clear()
    run1 = p1.add_run("▌ INSIGHT STRATEGICO")
    run1.bold = True
    run1.font.size = Pt(11)
    run1.font.color.rgb = BLU_SCURO
    
    # Testo insight
    p2 = cell2.add_paragraph()
    run2 = p2.add_run(text)
    run2.font.size = Pt(11)
    run2.font.color.rgb = NERO
    
    doc.add_paragraph()


def add_bullet_list(doc, items):
    """Aggiunge una lista puntata."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = NERO


def add_footer(doc):
    """Aggiunge footer a tutte le sezioni."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run("Evolution PRO  •  www.evolution-pro.it  •  Documento riservato")
        run.font.size = Pt(9)
        run.font.color.rgb = GRIGIO


# ═══════════════════════════════════════════════════════════════════════════════
# GENERAZIONE PAGINE
# ═══════════════════════════════════════════════════════════════════════════════

def create_cover_page(doc, dati):
    """Crea la copertina."""
    # Logo placeholder (testo centrato)
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EVOLUTION PRO")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BLU_SCURO
    
    doc.add_paragraph()
    
    # Titolo principale
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ANALISI STRATEGICA")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = BLU_SCURO
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PERSONALIZZATA")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = BLU_SCURO
    
    doc.add_paragraph()
    
    # Sottotitolo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Trasformare la tua competenza")
    run.font.size = Pt(16)
    run.font.color.rgb = GRIGIO
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("in un Asset Digitale")
    run.font.size = Pt(16)
    run.font.color.rgb = GRIGIO
    
    # Linea gialla centrata
    for _ in range(2):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 20)
    run.font.color.rgb = GIALLO
    
    for _ in range(2):
        doc.add_paragraph()
    
    # Info cliente
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Cliente")
    run.font.size = Pt(12)
    run.font.color.rgb = GRIGIO
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(dati.get("NOME_CLIENTE", ""))
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = BLU_SCURO
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ambito")
    run.font.size = Pt(12)
    run.font.color.rgb = GRIGIO
    
    # Estrai ambito dalla risposta expertise (prime 50 parole)
    ambito = dati.get("AMBITO", "")[:100]
    if len(ambito) > 50:
        ambito = ambito[:50] + "..."
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ambito)
    run.font.size = Pt(14)
    run.font.color.rgb = BLU_SCURO
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Data")
    run.font.size = Pt(12)
    run.font.color.rgb = GRIGIO
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(dati.get("DATA_ANALISI", ""))
    run.font.size = Pt(14)
    run.font.color.rgb = BLU_SCURO


def create_intro_page(doc):
    """Pagina introduttiva."""
    add_section_title(doc, "COSA TROVERAI IN QUESTA ANALISI")
    
    add_body_text(doc, "Questo documento è il risultato dell'analisi delle informazioni che hai condiviso nel questionario iniziale.")
    add_body_text(doc, "Il suo obiettivo è semplice: capire se la tua competenza può essere trasformata in un asset digitale sostenibile nel tempo.")
    
    doc.add_paragraph()
    add_body_text(doc, "In queste pagine troverai:")
    
    add_bullet_list(doc, [
        "una lettura strategica della tua situazione attuale",
        "un'analisi del mercato e del tuo pubblico",
        "l'individuazione dell'asset digitale più adatto",
        "una possibile evoluzione del tuo modello di lavoro",
        "una valutazione realistica della fattibilità del progetto"
    ])
    
    doc.add_paragraph()
    add_body_text(doc, "Questa analisi non nasce per convincerti.")
    add_body_text(doc, "Nasce per aiutarti a prendere una decisione lucida e consapevole sul prossimo passo del tuo percorso professionale.")
    
    doc.add_paragraph()
    add_body_text(doc, "Durante la videocall strategica commenteremo insieme questo documento punto per punto.")
    
    # Chi sono
    doc.add_paragraph()
    doc.add_paragraph()
    add_subtitle(doc, "Chi sono")
    
    add_body_text(doc, "Mi chiamo Claudio Bertogliatti.")
    add_body_text(doc, "Sono il fondatore di Evolution PRO, una digital agency specializzata nella creazione e vendita di videocorsi online per liberi professionisti.")
    add_body_text(doc, "Il nostro lavoro consiste nell'aiutare professionisti, coach, consulenti e formatori a trasformare la propria competenza in un asset digitale scalabile.")


def create_section_page(doc, section_num, title, label_risposta, risposta, considerazioni, insight=None):
    """Crea una pagina di sezione standard."""
    add_section_title(doc, f"SEZIONE {section_num} — {title}")
    
    add_client_response_box(doc, label_risposta, risposta)
    
    add_subtitle(doc, "Considerazioni strategiche")
    add_body_text(doc, considerazioni)
    
    if insight:
        doc.add_paragraph()
        add_insight_box(doc, insight)


def create_fai_da_te_page(doc):
    """Pagina 'Perché il fai-da-te fallisce'."""
    add_section_title(doc, "PERCHÉ IL FAI-DA-TE FALLISCE NEL 90% DEI CASI")
    
    add_body_text(doc, "Molti professionisti pensano che basti:")
    add_bullet_list(doc, [
        "registrare qualche video",
        "caricarlo online",
        "pubblicare qualche post sui social"
    ])
    
    add_body_text(doc, "Nella realtà quasi sempre accade il contrario.")
    add_body_text(doc, "Il corso viene registrato ma non viene venduto.")
    
    doc.add_paragraph()
    add_body_text(doc, "Questo accade perché un videocorso non è semplicemente un contenuto.")
    add_body_text(doc, "È un sistema.")
    
    doc.add_paragraph()
    add_body_text(doc, "Un sistema che include:")
    add_bullet_list(doc, [
        "posizionamento",
        "struttura del percorso",
        "asset di vendita",
        "funnel",
        "automazioni",
        "strategia di traffico"
    ])
    
    doc.add_paragraph()
    add_insight_box(doc, "Se uno di questi elementi manca, il sistema non funziona.")


def create_tre_strade_page(doc):
    """Pagina 'Le tre strade possibili'."""
    add_section_title(doc, "LE TRE STRADE POSSIBILI DA QUI")
    
    add_subtitle(doc, "1. NON FARE NULLA")
    add_body_text(doc, "Continuare con il modello attuale. Funziona, ma resta legato al tempo disponibile.")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("─" * 30)
    run.font.color.rgb = RGBColor(0xE5, 0xE7, 0xEB)
    
    add_subtitle(doc, "2. PROVARE DA SOLO")
    add_body_text(doc, "È una strada possibile, ma spesso porta a mesi di tentativi senza una struttura chiara.")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("─" * 30)
    run.font.color.rgb = RGBColor(0xE5, 0xE7, 0xEB)
    
    add_subtitle(doc, "3. COSTRUIRE UNA STRUTTURA")
    add_body_text(doc, "Lavorare con un metodo già testato e un team che conosce gli errori da evitare.")
    
    doc.add_paragraph()
    add_insight_box(doc, "Ed è esattamente il motivo per cui esiste Evolution PRO.")


def create_diagnosi_page(doc, fattibilita="ALTA"):
    """Pagina diagnosi finale."""
    add_section_title(doc, "DIAGNOSI STRATEGICA")
    
    add_body_text(doc, "Dall'analisi delle informazioni condivise emerge che la tua competenza può essere strutturata in un percorso digitale trasferibile.")
    add_body_text(doc, "Il modello attuale con cui viene proposta appare ancora fortemente legato al tempo disponibile.")
    add_body_text(doc, "La creazione di un asset digitale può rappresentare una evoluzione naturale del tuo percorso professionale.")
    add_body_text(doc, "Non come sostituzione del lavoro attuale, ma come estensione del tuo metodo.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Box fattibilità
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("La fattibilità del progetto appare:")
    run.font.size = Pt(14)
    run.font.color.rgb = GRIGIO
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(fattibilita)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BLU_SCURO
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_body_text(doc, "Il prossimo passo sarà la videocall strategica in cui analizzeremo insieme questa diagnosi e valuteremo la possibile evoluzione del progetto.")
    
    # Footer finale
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 20)
    run.font.color.rgb = GIALLO
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Evolution PRO")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLU_SCURO
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Trasformiamo competenze reali in asset digitali sostenibili.")
    run.font.size = Pt(11)
    run.font.color.rgb = GRIGIO
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("www.evolution-pro.it")
    run.font.size = Pt(11)
    run.font.color.rgb = GRIGIO


# ═══════════════════════════════════════════════════════════════════════════════
# FUNZIONE PRINCIPALE
# ═══════════════════════════════════════════════════════════════════════════════

def genera_analisi_strategica(dati: dict, output_path: str) -> str:
    """
    Genera il documento DOCX dell'Analisi Strategica.
    
    Args:
        dati: Dizionario con tutti i dati (risposte + AI)
        output_path: Percorso del file di output
    
    Returns:
        Percorso del file generato
    """
    doc = Document()
    
    # Imposta margini
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 1: COPERTINA
    # ═══════════════════════════════════════════════════════════════════════════
    create_cover_page(doc, dati)
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 2: INTRODUZIONE
    # ═══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    create_intro_page(doc)
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 3: ANALISI PROFILO PROFESSIONALE
    # ═══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    create_section_page(
        doc,
        section_num=1,
        title="ANALISI DEL TUO PROFILO PROFESSIONALE",
        label_risposta="Tua risposta: In cosa sei riconosciuto come esperto",
        risposta=dati.get("RISPOSTA_EXPERTISE", ""),
        considerazioni=dati.get("COMPETENZA_ARRICCHITA", "Analisi in elaborazione..."),
        insight=None
    )
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 4: CLIENTE IDEALE
    # ═══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    create_section_page(
        doc,
        section_num=2,
        title="IL TUO CLIENTE IDEALE",
        label_risposta="Tua risposta: Target dichiarato",
        risposta=dati.get("RISPOSTA_TARGET", ""),
        considerazioni=dati.get("TARGET_ARRICCHITO", "Analisi in elaborazione...")
    )
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 5: PRESENZA ONLINE
    # ═══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    create_section_page(
        doc,
        section_num=3,
        title="PRESENZA ONLINE",
        label_risposta="Tua risposta: Community e pubblico esistente",
        risposta=dati.get("RISPOSTA_PUBBLICO", ""),
        considerazioni=dati.get("PUBBLICO_ARRICCHITO", "Analisi in elaborazione...")
    )
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 6: ESPERIENZA DI VENDITA
    # ═══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    create_section_page(
        doc,
        section_num=4,
        title="ESPERIENZA DI VENDITA ONLINE",
        label_risposta="Tua risposta: Esperienze precedenti",
        risposta=dati.get("RISPOSTA_ESPERIENZA", ""),
        considerazioni=dati.get("ESPERIENZA_ARRICCHITA", "Analisi in elaborazione...")
    )
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 7: OSTACOLO PRINCIPALE
    # ═══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    create_section_page(
        doc,
        section_num=5,
        title="L'OSTACOLO PRINCIPALE",
        label_risposta="Tua risposta: Il blocco che hai incontrato finora",
        risposta=dati.get("RISPOSTA_OSTACOLO", ""),
        considerazioni=dati.get("OSTACOLO_ARRICCHITO", "Analisi in elaborazione...")
    )
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 8: OBIETTIVO / RISULTATO CONCRETO
    # ═══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    create_section_page(
        doc,
        section_num=6,
        title="IL RISULTATO CHE VUOI OTTENERE",
        label_risposta="Tua risposta: Obiettivo per i prossimi 12 mesi",
        risposta=dati.get("RISPOSTA_OBIETTIVO", dati.get("RISPOSTA_RISULTATO", "")),
        considerazioni=dati.get("OBIETTIVO_ARRICCHITO", dati.get("RISULTATO_ARRICCHITO", "Analisi in elaborazione..."))
    )
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 9: PERCHÉ ADESSO
    # ═══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    create_section_page(
        doc,
        section_num=7,
        title="PERCHÉ PROPRIO ADESSO",
        label_risposta="Tua risposta",
        risposta=dati.get("RISPOSTA_PERCHE_ORA", ""),
        considerazioni=dati.get("PERCHE_ORA_ARRICCHITO", "Analisi in elaborazione...")
    )
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 10: PERCHÉ IL FAI-DA-TE FALLISCE
    # ═══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    create_fai_da_te_page(doc)
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 11: LE TRE STRADE
    # ═══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    create_tre_strade_page(doc)
    
    # ═══════════════════════════════════════════════════════════════════════════
    # PAGINA 12: DIAGNOSI FINALE
    # ═══════════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    fattibilita = dati.get("FATTIBILITA", "ALTA")
    create_diagnosi_page(doc, fattibilita)
    
    # Aggiungi footer
    add_footer(doc)
    
    # Salva documento
    doc.save(output_path)
    
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Uso: python genera_analisi_docx.py '<json_dati>' '<output_path>'")
        sys.exit(1)
    
    try:
        dati = json.loads(sys.argv[1])
        output_path = sys.argv[2]
        
        result = genera_analisi_strategica(dati, output_path)
        print(f"OK: {result}")
        
    except Exception as e:
        print(f"ERRORE: {e}")
        sys.exit(1)
